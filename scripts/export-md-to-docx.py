#!/usr/bin/env python3
"""Export a Markdown SRS to DOCX using a reference DOCX template."""

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path


DEFAULT_MD = Path.home() / "Downloads" / "hotel_management_system_srs_v1_2_staff_screen_mockup_ready.md"
DEFAULT_TEMPLATE_NAME = "26_01_2026___3324605acbaa0a423a9bc9ef796329a2. (2).docx"
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def parse_args() -> argparse.Namespace:
    script_dir = Path(__file__).resolve().parent
    default_template = script_dir.parent / DEFAULT_TEMPLATE_NAME
    parser = argparse.ArgumentParser(description="Export Markdown to DOCX with a reference template.")
    parser.add_argument("--input", "-i", type=Path, default=DEFAULT_MD, help="Markdown input path.")
    parser.add_argument("--template", "-t", type=Path, default=default_template, help="Reference DOCX template.")
    parser.add_argument("--output", "-o", type=Path, help="Output DOCX path.")
    parser.add_argument("--engine", choices=["auto", "pandoc", "python-docx"], default="auto")
    parser.add_argument("--resource-path", action="append", type=Path, default=[], help="Extra image/resource root.")
    parser.add_argument("--toc", action="store_true", help="Add a table of contents when using Pandoc.")
    parser.add_argument("--overwrite", action="store_true", help="Overwrite output if it already exists.")
    parser.add_argument(
        "--mockup-width-inches",
        type=float,
        default=2.65,
        help="Rendered width for mobile mock-up images in DOCX.",
    )
    parser.add_argument(
        "--image-width-inches",
        type=float,
        default=6.2,
        help="Rendered width for non-mock-up images in DOCX.",
    )
    return parser.parse_args()


def normalize_path(path: Path) -> Path:
    return path.expanduser().resolve()


def ensure_paths(args: argparse.Namespace) -> tuple[Path, Path, Path]:
    md = normalize_path(args.input)
    template = normalize_path(args.template)
    output = normalize_path(args.output) if args.output else template.parent / f"{md.stem}.docx"
    if not md.exists():
        raise FileNotFoundError(f"Markdown input not found: {md}")
    if not template.exists():
        raise FileNotFoundError(f"Template DOCX not found: {template}")
    if output.exists() and not args.overwrite:
        raise FileExistsError(f"Output exists; pass --overwrite: {output}")
    output.parent.mkdir(parents=True, exist_ok=True)
    return md, template, output


def run_pandoc(md: Path, template: Path, output: Path, args: argparse.Namespace) -> None:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise RuntimeError("Pandoc is not installed or not in PATH.")
    roots = [md.parent, template.parent, *[normalize_path(p) for p in args.resource_path]]
    cmd = [
        pandoc,
        str(md),
        "--from",
        "gfm+pipe_tables+fenced_code_blocks+tex_math_dollars",
        "--to",
        "docx",
        "--standalone",
        "--reference-doc",
        str(template),
        "--resource-path",
        os.pathsep.join(str(p) for p in roots),
        "--output",
        str(output),
    ]
    if args.toc:
        cmd.append("--toc")
    subprocess.run(cmd, check=True)


def import_docx():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for --engine python-docx. Install with: "
            "python -m pip install --user python-docx"
        ) from exc
    return Document, WD_ALIGN_PARAGRAPH, qn, Pt, Inches


def clear_body(document, qn) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def style_name(document, *candidates: str) -> str | None:
    by_name = {s.name.lower(): s.name for s in document.styles}
    by_id = {getattr(s, "style_id", "").lower(): s.name for s in document.styles}
    for item in candidates:
        key = item.lower()
        if key in by_name:
            return by_name[key]
        if key in by_id:
            return by_id[key]
    return None


def add_runs(paragraph, text: str) -> None:
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    for part in token_re.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(part)


def set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text_color(cell, color: str) -> None:
    from docx.shared import RGBColor

    rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = rgb


def set_cell_bold(cell, bold: bool = True) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = bold


def apply_generic_table_style(table) -> None:
    for r_idx, row in enumerate(table.rows):
        fill = "FFFFFF"
        if r_idx == 0:
            fill = "D9EAF7"
        elif r_idx % 2 == 0:
            fill = "F7FAFC"
        for cell in row.cells:
            set_cell_shading(cell, fill)
            if r_idx == 0:
                set_cell_bold(cell)
                set_cell_text_color(cell, "17365D")


def apply_use_case_table_style(table, table_kind: str) -> None:
    palettes = {
        "main": ("1F4E79", "D9EAF7", "EEF5FB"),
        "alternative": ("604A7B", "E4DFEC", "F7F4FA"),
        "business": ("548235", "E2F0D9", "F6FAF3"),
    }
    dark, label, zebra = palettes.get(table_kind, palettes["main"])
    for r_idx, row in enumerate(table.rows):
        row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        unique_values = set(row_values)
        is_section_row = len(unique_values) == 1 and next(iter(unique_values), "") in {
            "Main flows",
            "Alternative flows",
            "Business Rules",
        }
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if r_idx == 0 or is_section_row:
                set_cell_shading(cell, dark)
                set_cell_text_color(cell, "FFFFFF")
                set_cell_bold(cell)
            elif text in {
                "Use Case ID",
                "Use Case Name",
                "Author",
                "Version",
                "Date",
                "Actor",
                "Secondary Actor(s)",
                "Feature / Group Function",
                "Description",
                "Precondition",
                "Trigger",
                "Post-Condition",
                "Step",
                "Sub step",
                "Action",
                "Alternative ID",
                "Condition",
                "#",
                "Rule Description",
            }:
                set_cell_shading(cell, label)
                set_cell_text_color(cell, "17365D")
                set_cell_bold(cell)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, zebra)


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.replace("\\|", "|").strip() for cell in line.split("|")]


def resolve_resource(target: str, roots: list[Path]) -> Path | None:
    raw = target.strip().strip("<>")
    path = Path(raw)
    if path.is_absolute() and path.exists():
        return path
    for root in roots:
        candidate = (root / raw).resolve()
        if candidate.exists():
            return candidate
    return None


def add_table(document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    style = style_name(document, "Table Grid", "TableGrid", "Kiu2")
    if style:
        table.style = style
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_runs(paragraph, text)
            if r_idx == 0:
                for run in paragraph.runs:
                    run.bold = True
    apply_generic_table_style(table)


def add_html_table(document, html_block: str) -> None:
    try:
        from lxml import html
    except ImportError as exc:
        raise RuntimeError("lxml is required to export HTML tables.") from exc
    root = html.fragment_fromstring(html_block, create_parent="div")
    table_el = root.find(".//table")
    if table_el is None:
        return
    rows = table_el.findall(".//tr")
    if not rows:
        return
    grid: list[list[dict[str, object] | None]] = []
    max_cols = 0
    for tr in rows:
        row: list[dict[str, object] | None] = []
        col_idx = 0
        for cell_el in tr.xpath("./th|./td"):
            while col_idx < len(row) and row[col_idx] is not None:
                col_idx += 1
            colspan = int(cell_el.get("colspan", "1") or "1")
            text = " ".join(cell_el.text_content().split())
            row.extend([None] * max(0, col_idx - len(row)))
            row.append({"text": text, "colspan": colspan, "header": cell_el.tag.lower() == "th"})
            for _ in range(colspan - 1):
                row.append(None)
            col_idx += colspan
        max_cols = max(max_cols, len(row))
        grid.append(row)
    docx_table = document.add_table(rows=len(grid), cols=max_cols)
    style = style_name(document, "Table Grid", "TableGrid", "Kiu2")
    if style:
        docx_table.style = style
    for r_idx, row in enumerate(grid):
        c_idx = 0
        while c_idx < max_cols:
            item = row[c_idx] if c_idx < len(row) else None
            if item is None:
                c_idx += 1
                continue
            cell = docx_table.cell(r_idx, c_idx)
            colspan = int(item["colspan"])
            if colspan > 1 and c_idx + colspan - 1 < max_cols:
                cell = cell.merge(docx_table.cell(r_idx, c_idx + colspan - 1))
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_runs(paragraph, str(item["text"]))
            if item["header"]:
                for run in paragraph.runs:
                    run.bold = True
            c_idx += colspan
    first_text = " ".join(cell.text.strip() for cell in docx_table.rows[0].cells if cell.text.strip())
    if "Use Case ID" in first_text:
        apply_use_case_table_style(docx_table, "main")
    elif "Alternative flows" in first_text:
        apply_use_case_table_style(docx_table, "alternative")
    elif "Business Rules" in first_text:
        apply_use_case_table_style(docx_table, "business")
    else:
        apply_generic_table_style(docx_table)


def add_image(document, image_path: Path, alt: str, WD_ALIGN_PARAGRAPH, Inches, args: argparse.Namespace) -> None:
    section = document.sections[-1]
    max_width = section.page_width - section.left_margin - section.right_margin
    requested_width = args.image_width_inches
    if image_path.name.startswith("MCK-SCR-") or "mobile-flutter" in image_path.name:
        requested_width = args.mockup_width_inches
    width = min(Inches(requested_width), max_width)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    try:
        run.add_picture(str(image_path), width=width)
    except Exception:
        paragraph.add_run(f"[Image not inserted: {alt or image_path.name}]")


def export_with_python_docx(md: Path, template: Path, output: Path, args: argparse.Namespace) -> None:
    Document, WD_ALIGN_PARAGRAPH, qn, Pt, Inches = import_docx()
    document = Document(str(template))
    clear_body(document, qn)
    roots = [md.parent, template.parent, *[normalize_path(p) for p in args.resource_path]]
    lines = md.read_text(encoding="utf-8-sig").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i].rstrip()
        if line.strip().startswith("```"):
            if in_code:
                para = document.add_paragraph()
                for idx, code_line in enumerate(code_lines):
                    if idx:
                        para.add_run("\n")
                    run = para.add_run(code_line)
                    run.font.name = "Courier New"
                    run.font.size = Pt(8)
                code_lines.clear()
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        if line.lstrip().lower().startswith("<table"):
            html_lines = [line]
            i += 1
            while i < len(lines):
                html_lines.append(lines[i])
                if "</table>" in lines[i].lower():
                    i += 1
                    break
                i += 1
            add_html_table(document, "\n".join(html_lines))
            continue
        if i + 1 < len(lines) and "|" in line and TABLE_SEP_RE.match(lines[i + 1]):
            table_rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                table_rows.append(split_table_row(lines[i]))
                i += 1
            add_table(document, table_rows)
            continue
        image_match = IMAGE_RE.fullmatch(line.strip())
        if image_match:
            alt, target = image_match.groups()
            resolved = resolve_resource(target, roots)
            if resolved:
                add_image(document, resolved, alt, WD_ALIGN_PARAGRAPH, Inches, args)
            else:
                document.add_paragraph(f"[Missing image: {target}]")
            i += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 5)
            style = style_name(document, f"Heading {level}", f"heading {level}", f"Heading{level}")
            para = document.add_paragraph(style=style) if style else document.add_paragraph()
            add_runs(para, heading.group(2).strip())
            i += 1
            continue
        caption = re.match(r"^(Figure|Table)\s+\d+-\d+:.+$", line.strip())
        if caption:
            style = style_name(document, "Caption", "caption")
            para = document.add_paragraph(style=style) if style else document.add_paragraph()
            add_runs(para, line.strip())
            i += 1
            continue
        list_item = re.match(r"^\s*(?:[-*+]|\d+\.)\s+(.+)$", line)
        if list_item:
            style = style_name(document, "List Paragraph", "ListParagraph", "List")
            para = document.add_paragraph(style=style) if style else document.add_paragraph()
            add_runs(para, f"- {list_item.group(1).strip()}")
            i += 1
            continue
        para = document.add_paragraph()
        add_runs(para, line.lstrip("> ").strip())
        i += 1
    document.save(str(output))


def main() -> int:
    args = parse_args()
    md, template, output = ensure_paths(args)
    engine = args.engine
    if engine == "auto":
        engine = "pandoc" if shutil.which("pandoc") else "python-docx"
    if engine == "pandoc":
        run_pandoc(md, template, output, args)
    else:
        export_with_python_docx(md, template, output, args)
    print(f"Exported DOCX: {output}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise SystemExit(1)
